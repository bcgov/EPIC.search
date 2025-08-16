"""
Word document reader for extracting text from DOCX and DOC files.

This module provides text extraction capabilities for Microsoft Word documents:
- DOCX files: Uses python-docx for structured text extraction
- DOC files: Uses docx2txt as a fallback for legacy formats
- Handles tables, headers, footers, and formatted text
- Preserves document structure where possible
"""

import os
import tempfile
from typing import List, Dict, Any
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False


def is_word_supported():
    """Check if Word document processing is available."""
    return DOCX_AVAILABLE or DOCX2TXT_AVAILABLE


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If extraction fails
    """
    if not DOCX_AVAILABLE:
        raise Exception("python-docx not available for DOCX processing")
    
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # Join all content
        full_text = "\n\n".join(text_content)
        
        if not full_text.strip():
            raise Exception("No text content extracted from DOCX file")
            
        return full_text
        
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_doc_fallback(file_path: str) -> str:
    """
    Extract text from DOC/DOCX files using docx2txt as fallback.
    
    Args:
        file_path: Path to the DOC/DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If extraction fails
    """
    if not DOCX2TXT_AVAILABLE:
        raise Exception("docx2txt not available for DOC/DOCX processing")
    
    try:
        # docx2txt can handle both DOC and DOCX files
        text_content = docx2txt.process(file_path)
        
        if not text_content or not text_content.strip():
            raise Exception("No text content extracted from document")
            
        return text_content.strip()
        
    except Exception as e:
        raise Exception(f"Failed to extract text using docx2txt: {str(e)}")


def extract_text_from_word_document(file_path: str) -> str:
    """
    Extract text from Word documents (DOCX or DOC).
    
    This function tries multiple extraction methods:
    1. python-docx for DOCX files (preferred for structure)
    2. docx2txt as fallback for both DOC and DOCX
    
    Args:
        file_path: Path to the Word document
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If all extraction methods fail
    """
    file_ext = Path(file_path).suffix.lower()
    
    # Try python-docx first for DOCX files
    if file_ext == '.docx' and DOCX_AVAILABLE:
        try:
            return extract_text_from_docx(file_path)
        except Exception as e:
            print(f"[WARNING] python-docx failed for {file_path}: {e}")
            # Fall back to docx2txt
    
    # Try docx2txt for both DOC and DOCX
    if DOCX2TXT_AVAILABLE:
        try:
            return extract_text_from_doc_fallback(file_path)
        except Exception as e:
            print(f"[WARNING] docx2txt failed for {file_path}: {e}")
    
    # If we get here, all methods failed
    available_methods = []
    if DOCX_AVAILABLE:
        available_methods.append("python-docx")
    if DOCX2TXT_AVAILABLE:
        available_methods.append("docx2txt")
    
    if not available_methods:
        raise Exception("No Word document processing libraries available. Install python-docx or docx2txt.")
    else:
        raise Exception(f"All available extraction methods failed: {', '.join(available_methods)}")


def read_word_as_pages(file_path: str) -> List[Dict[str, Any]]:
    """
    Read a Word document and return it as page-like chunks for compatibility with PDF processing.
    
    Since Word documents don't have fixed pages like PDFs, we simulate pages by:
    - Splitting content into reasonable chunks
    - Maintaining compatibility with existing PDF processing pipeline
    
    Args:
        file_path: Path to the Word document
        
    Returns:
        List of page dictionaries with 'content' and 'page_number' keys
    """
    try:
        # Extract text content
        full_text = extract_text_from_word_document(file_path)
        
        # Split into logical "pages" (chunks of ~2000 characters to simulate pages)
        chunk_size = 2000
        text_chunks = []
        
        # Split by double newlines first (paragraphs)
        paragraphs = full_text.split('\n\n')
        current_chunk = ""
        
        for paragraph in paragraphs:
            # If adding this paragraph would exceed chunk size, start new chunk
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                text_chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk += ("\n\n" if current_chunk else "") + paragraph
        
        # Add the last chunk
        if current_chunk.strip():
            text_chunks.append(current_chunk.strip())
        
        # Convert to page format
        pages = []
        for i, chunk in enumerate(text_chunks, 1):
            pages.append({
                'content': chunk,
                'page_number': i,
                'extraction_method': 'word_document'
            })
        
        if not pages:
            # Fallback: treat entire document as one page
            pages = [{
                'content': full_text,
                'page_number': 1,
                'extraction_method': 'word_document'
            }]
        
        return pages
        
    except Exception as e:
        raise Exception(f"Failed to read Word document {file_path}: {str(e)}")


def get_word_document_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from Word documents.
    
    Args:
        file_path: Path to the Word document
        
    Returns:
        Dictionary containing document metadata
    """
    metadata = {
        'file_type': 'word_document',
        'file_extension': Path(file_path).suffix.lower(),
        'extraction_methods_available': []
    }
    
    if DOCX_AVAILABLE:
        metadata['extraction_methods_available'].append('python-docx')
    if DOCX2TXT_AVAILABLE:
        metadata['extraction_methods_available'].append('docx2txt')
    
    # Try to get document properties for DOCX files
    if Path(file_path).suffix.lower() == '.docx' and DOCX_AVAILABLE:
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata.update({
                'title': core_props.title or 'Unknown',
                'author': core_props.author or 'Unknown',
                'created': str(core_props.created) if core_props.created else 'Unknown',
                'modified': str(core_props.modified) if core_props.modified else 'Unknown',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            })
        except Exception as e:
            print(f"[WARNING] Could not extract metadata from {file_path}: {e}")
            metadata.update({
                'title': 'Unknown',
                'author': 'Unknown',
                'created': 'Unknown',
                'modified': 'Unknown'
            })
    else:
        # Basic metadata for DOC files or when python-docx is not available
        metadata.update({
            'title': Path(file_path).stem,
            'author': 'Unknown',
            'created': 'Unknown',
            'modified': 'Unknown'
        })
    
    return metadata
